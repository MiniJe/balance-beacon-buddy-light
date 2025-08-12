#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Versiune LIGHT a generatorului avansat de PDF-uri pentru Balance Beacon Buddy
Adaptat pentru:
 - SQLite local (în loc de Azure SQL / API remote)
 - Template-uri DOCX locale (cu diacritice)
 - Generare PDF locală folosind python-docx + docx2pdf (sau fallback la conversie simplă)

Expected usage (silent JSON mode):
  python advanced_pdf_generator.py \
      --partner-id <ID> \
      --db-path <cale_sqlite> \
      --nr-document <numar> \
      --data-emiterii DD.MM.YYYY \
      --data-sold DD.MM.YYYY \
      --template-name document_template_clienți-duc.docx \
      --template-path C:/BalanceBeaconBuddy/Șabloane/document_template_clienți-duc.docx \
      --output-dir C:/CereriConfirmare/ \
      --json

Output JSON keys:
  success, error (optional), nr_document, template_used, docx_path, pdf_path

IMPORTANT: Dacă nu poate genera PDF (lipsește docx2pdf), va livra DOCX și setează pdf_path la None.
"""
import argparse
import json
import os
import sys
import sqlite3
import datetime
import hashlib
import tempfile
import re

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from docx2pdf import convert as docx2pdf_convert
except ImportError:
    docx2pdf_convert = None

# Setul EXACT de placeholdere din sabloane (fara aliasuri suplimentare)
PLACEHOLDERS_MAP = {
    '{COMPANIE}': 'numePartener',
    '{C.U.I.}': 'cuiPartener',
    '{O.N.R.C.}': 'onrcPartener',
    '{DATA-EMITERII}': 'dataEmiterii',
    '{DATA-SOLD}': 'dataSold',
    '{NR.DOC.}': 'nrDocument'
}
# Regex generic care permite orice caracter (in afara de acolade) intre { }
PLACEHOLDER_TOKEN_REGEX = re.compile(r'\{[^{}]+\}')

def read_partner(conn, partner_id: str):
    cur = conn.cursor()
    cur.execute("SELECT IdPartener, NumePartener, CUIPartener, ONRCPartener, EmailPartener, ReprezentantPartener FROM Parteneri WHERE IdPartener = ? LIMIT 1", (partner_id,))
    row = cur.fetchone()
    if not row:
        raise ValueError(f'Partenerul cu ID {partner_id} nu a fost găsit')
    return {
        'idPartener': row[0],
        'numePartener': row[1] or '',
        'cuiPartener': row[2] or '',
        'onrcPartener': row[3] or '',
        'emailPartener': row[4] or '',
        'reprezentantPartener': row[5] or ''
    }

def replace_placeholders(doc, context: dict, debug: bool = False):
    def build_value(ph: str) -> str:
        key = PLACEHOLDERS_MAP.get(ph)
        return '' if key is None else str(context.get(key, '') or '')

    def process_paragraph(paragraph):
        if not paragraph.runs:
            return
        original_runs = [r.text for r in paragraph.runs]
        full_text = ''.join(original_runs)
        changed = False
        # Inlocuire directa pentru fiecare placeholder definit
        for ph in PLACEHOLDERS_MAP.keys():
            if ph in full_text:
                full_text = full_text.replace(ph, build_value(ph))
                changed = True
        if changed:
            paragraph.runs[0].text = full_text
            for r in paragraph.runs[1:]:
                r.text = ''

    for p in doc.paragraphs:
        process_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_paragraph(p)
    if debug:
        remaining = []
        for p in doc.paragraphs:
            remaining.extend([tok for tok in PLACEHOLDER_TOKEN_REGEX.findall(p.text) if tok in PLACEHOLDERS_MAP])
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        remaining.extend([tok for tok in PLACEHOLDER_TOKEN_REGEX.findall(p.text) if tok in PLACEHOLDERS_MAP])
        if remaining:
            print('[DEBUG] Placeholdere neînlocuite:', sorted(set(remaining)))

def generate_docx(template_path: str, output_dir: str, output_name: str, context: dict) -> str:
    if not Document:
        raise RuntimeError('Lipsește pachetul python-docx. Instalați: pip install python-docx')
    doc = Document(template_path)
    replace_placeholders(doc, context)
    out_path = os.path.join(output_dir, output_name)
    doc.save(out_path)
    return out_path

def try_convert_pdf(docx_path: str) -> str | None:
    if not docx2pdf_convert:
        return None
    # Creează un folder temporar pentru conversie (docx2pdf cere mediu controlat uneori)
    temp_dir = tempfile.mkdtemp(prefix='bbbuddy_pdf_')
    try:
        candidate_pdf = os.path.splitext(os.path.basename(docx_path))[0] + '.pdf'
        temp_pdf_path = os.path.join(temp_dir, candidate_pdf)
        # docx2pdf convert in-place folder->folder
        docx2pdf_convert(docx_path, temp_pdf_path)
        if os.path.exists(temp_pdf_path):
            final_pdf_path = os.path.splitext(docx_path)[0] + '.pdf'
            try:
                os.replace(temp_pdf_path, final_pdf_path)
            except Exception:
                # fallback copy
                import shutil
                shutil.copy2(temp_pdf_path, final_pdf_path)
            return final_pdf_path
        return None
    except Exception:
        return None

def sha256_file(path_: str) -> str:
    h = hashlib.sha256()
    with open(path_, 'rb') as f:
        for chunk in iter(lambda: f.read(8192), b''):
            h.update(chunk)
    return h.hexdigest()

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--partner-id', required=True)
    parser.add_argument('--db-path', required=True)
    parser.add_argument('--nr-document', required=True, type=int)
    parser.add_argument('--data-emiterii', required=True)
    parser.add_argument('--data-sold', required=True)
    parser.add_argument('--template-name', required=True)
    parser.add_argument('--template-path', required=True)
    parser.add_argument('--output-dir', required=True)
    parser.add_argument('--json', action='store_true')
    parser.add_argument('--debug', action='store_true')
    args = parser.parse_args()
    try:
        if not os.path.exists(args.db_path):
            raise FileNotFoundError(f'Baza de date SQLite nu a fost găsită: {args.db_path}')
        if not os.path.exists(args.template_path):
            raise FileNotFoundError(f'Template-ul nu a fost găsit: {args.template_path}')
        os.makedirs(args.output_dir, exist_ok=True)
        conn = sqlite3.connect(args.db_path)
        try:
            partner = read_partner(conn, args.partner_id)
        finally:
            conn.close()
        if args.debug:
            print(f"[DEBUG] Partener: {partner['numePartener']} ({partner['idPartener']})")
            print(f"[DEBUG] Raw CUI={partner['cuiPartener']} ONRC={partner['onrcPartener']}")
        # Fara normalizare / swap – se folosesc valorile exact cum sunt in DB
        context = {
            'numePartener': partner['numePartener'],
            'cuiPartener': partner['cuiPartener'],
            'onrcPartener': partner['onrcPartener'],
            'dataSold': args.data_sold,
            'dataEmiterii': args.data_emiterii,
            'nrDocument': str(args.nr_document)
        }
        now = datetime.datetime.now()
        date_compact = args.data_sold.replace('.', '') if '.' in args.data_sold else args.data_sold.replace('-', '')
        time_part = now.strftime('%H%M%S')
        partener_name_clean = partner['numePartener'].replace('/', ' ').replace('\\', ' ').replace(':', ' ').replace('*', ' ').replace('?', ' ').replace('"', ' ').replace('<', ' ').replace('>', ' ').replace('|', ' ').strip()
        cui_clean = (partner['cuiPartener'] or 'FARA_CUI').replace('/', '').replace('\\', '').replace(':', '').replace('*', '').replace('?', '').replace('"', '').replace('<', '').replace('>', '').replace('|', '').strip()
        base_name = f"Nr{args.nr_document}_CERERE_SOLD_{partener_name_clean}_{cui_clean}_{date_compact}_{time_part}"
        docx_name = base_name + '.docx'
        docx_path = generate_docx(args.template_path, args.output_dir, docx_name, context)
        pdf_path = try_convert_pdf(docx_path)
        result = {
            'success': True,
            'nr_document': args.nr_document,
            'template_used': args.template_name,
            'docx_path': docx_path,
            'pdf_path': pdf_path,
        }
        print(json.dumps(result, ensure_ascii=False))
    except Exception as e:
        if '--json' in sys.argv:
            print(json.dumps({'success': False, 'error': str(e)}, ensure_ascii=False))
        else:
            print(f'Eroare: {e}', file=sys.stderr)
        sys.exit(1)

if __name__ == '__main__':
    main()
